import io
import os
import re
import json
import hashlib
import zipfile
import time
from pathlib import Path
from collections import Counter, defaultdict

import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from google import genai
from google.genai import types

APP_NAME = "Coach Winnie – Forms Converter"
APP_VERSION = "V6.10 Clear API Status"
DEFAULT_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.6-flash")

st.set_page_config(page_title=APP_NAME, page_icon="📝", layout="centered")

st.markdown("""
<style>
.block-container {max-width: 960px; padding-top: 2rem; padding-bottom: 3rem;}
.hero {
  padding: 1.4rem 1.6rem;
  border-radius: 22px;
  background: linear-gradient(135deg, rgba(98,72,255,.12), rgba(40,180,160,.10));
  border: 1px solid rgba(120,120,120,.18);
  margin-bottom: 1.1rem;
}
.hero h1 {margin:0; font-size:2rem;}
.hero p {margin:.45rem 0 0 0; opacity:.78;}
.small {font-size:.9rem; opacity:.72;}
</style>
""", unsafe_allow_html=True)

st.markdown(f"""
<div class="hero">
  <h1>📝 {APP_NAME}</h1>
  <p>Google Forms PDF → Microsoft Forms Quick Import Word (.docx)</p>
  <p><strong>{APP_VERSION}</strong></p>
</div>
""", unsafe_allow_html=True)

st.info(
    "V6.5 只使用 3 种题型：Choice、Multiple-answer Choice、Open text。"
    "Gemini 不再判断复杂的 question_type / original_type / conversion_action。"
    "图片提取仍独立进行，不影响题型分类。"
    "注意：Microsoft Forms Quick Import 可能不会把 Word 内图片自动带入题目。"
)

st.subheader("🔑 使用自己的 Gemini API Key")
st.markdown(
    "第一次使用？可到 **Google AI Studio** 建立 Gemini API Key："
    "[Get API Key](https://aistudio.google.com/app/apikey)"
)
api_key_input = st.text_input(
    "Gemini API Key",
    type="password",
    placeholder="AIza...",
    help="每次使用时输入。此 App 不会把你的 API Key 写入 GitHub、文件或数据库。"
)
st.caption(
    "🛡️ V6.5：继续使用 gemini-3.6-flash；遇到 503 / high demand 只自动重试一次。题型识别仅保留 3 类，减少 Gemini 判断负担。\n\n"
    "🔒 API Key 只用于本次页面会话。关闭/刷新页面后请重新输入。"
    "同一个有效的 Gemini API Key 可以重复使用。"
)

with st.expander("V6.10 Clear API Status 转换规则", expanded=False):
    st.markdown("""
**按照 Microsoft Forms Import Guidance：**

- 推荐题型：**Multiple choice**、**Open text**
- 每一题之间要有清楚分隔
- 内容要**垂直排列**
- Quick Import Word **不放图片 / figures**
- 复杂 Grid / Likert / Matching 会先拆成独立 Choice
- Checkboxes / Checkbox grid 会输出为 Choice；导入后请手动开启 **Multiple answers**
- Short answer / Paragraph / 无法安全转换的题型 → Open text
- 图片会另外保存到 ZIP，不影响 Quick Import Word
- Quiz 模式若提供答案文件：Choice / Multiple-answer Choice 会输出 **Answer: A** 或 **Answer: A, C**
- 答案必须来自上传的答案文件；缺少答案时不会推测
""")

pdf_file = st.file_uploader("① 上传 Google Forms PDF", type=["pdf"])
answer_file = st.file_uploader(
    "② 答案文件（可选）",
    type=["pdf", "txt", "docx"],
    help="如果答案不在原 PDF 中，可另上传答案版。没有答案文件也可以直接转换。"
)

col1, col2 = st.columns(2)
with col1:
    output_mode = st.radio(
        "输出类型",
        ["Form（无答案）", "Quiz（有答案时加入答案）"],
        index=0
    )
with col2:
    st.selectbox(
        "界面提示语言",
        ["保持原文件语言", "中文提示", "English prompts"],
        index=0
    )


def extract_pdf_text(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    chunks = []
    for i, page in enumerate(doc):
        txt = page.get_text("text")
        chunks.append(f"\n===== PAGE {i+1} =====\n{txt}")
    return "\n".join(chunks)


def clean_form_text_for_gemini(text: str) -> str:
    """
    Remove common Google Forms / print-view UI noise while preserving:
    - PAGE markers
    - original question text
    - original options
    - section/title content
    """
    if not text:
        return ""

    noise_patterns = [
        r"^\s*Clear selection\s*$",
        r"^\s*Required\s*$",
        r"^\s*\*\s*$",
        r"^\s*Never submit passwords through Google Forms.*$",
        r"^\s*This form was created inside.*$",
        r"^\s*Report Abuse\s*$",
        r"^\s*Google Forms\s*$",
        r"^\s*Page\s+\d+\s+of\s+\d+\s*$",
        r"^\s*Back\s*$",
        r"^\s*Next\s*$",
        r"^\s*Submit\s*$",
    ]

    cleaned_lines = []
    last_blank = False

    for raw in text.splitlines():
        line = raw.rstrip()

        if any(re.match(pat, line, flags=re.I) for pat in noise_patterns):
            continue

        # Keep PAGE markers exactly because image questions may need source_page.
        if line.startswith("===== PAGE "):
            cleaned_lines.append(line)
            last_blank = False
            continue

        # Collapse repeated blank lines.
        if not line.strip():
            if not last_blank:
                cleaned_lines.append("")
            last_blank = True
            continue

        cleaned_lines.append(line)
        last_blank = False

    return "\n".join(cleaned_lines).strip()



def extract_pdf_images(data: bytes):
    """Extract useful raster images from PDF, keeping page metadata."""
    doc = fitz.open(stream=data, filetype="pdf")
    images = []
    seen = set()

    for page_no, page in enumerate(doc, start=1):
        page_imgs = page.get_images(full=True)
        page_index = 0

        for img in page_imgs:
            xref = img[0]
            try:
                info = doc.extract_image(xref)
            except Exception:
                continue

            blob = info.get("image", b"")
            width = int(info.get("width", 0) or 0)
            height = int(info.get("height", 0) or 0)
            ext = (info.get("ext") or "png").lower()

            # Filter tiny icons, bullets and low-value decoration.
            if not blob or len(blob) < 2500:
                continue
            if width < 120 or height < 80 or width * height < 25000:
                continue

            digest = hashlib.sha1(blob).hexdigest()
            if digest in seen:
                continue
            seen.add(digest)

            page_index += 1
            image_id = f"P{page_no:02d}_IMG{page_index:02d}"
            filename = f"page_{page_no:02d}_image_{page_index:02d}.{ext}"

            images.append({
                "id": image_id,
                "page": page_no,
                "index": page_index,
                "filename": filename,
                "ext": ext,
                "width": width,
                "height": height,
                "bytes": blob,
            })

    return images


def extract_pdf_images_for_pages(data: bytes, wanted_pages=None):
    """
    Lazy image extraction:
    - if wanted_pages is empty/None -> return []
    - otherwise extract only from those PDF pages
    """
    pages = {int(p) for p in (wanted_pages or []) if p}
    if not pages:
        return []

    doc = fitz.open(stream=data, filetype="pdf")
    images = []
    seen = set()

    for page_no, page in enumerate(doc, start=1):
        if page_no not in pages:
            continue

        page_imgs = page.get_images(full=True)
        page_index = 0

        for img in page_imgs:
            xref = img[0]
            try:
                info = doc.extract_image(xref)
            except Exception:
                continue

            blob = info.get("image", b"")
            width = int(info.get("width", 0) or 0)
            height = int(info.get("height", 0) or 0)
            ext = (info.get("ext") or "png").lower()

            if not blob or len(blob) < 2500:
                continue
            if width < 120 or height < 80 or width * height < 25000:
                continue

            digest = hashlib.sha1(blob).hexdigest()
            if digest in seen:
                continue
            seen.add(digest)

            page_index += 1
            image_id = f"P{page_no:02d}_IMG{page_index:02d}"
            filename = f"page_{page_no:02d}_image_{page_index:02d}.{ext}"

            images.append({
                "id": image_id,
                "page": page_no,
                "index": page_index,
                "filename": filename,
                "ext": ext,
                "width": width,
                "height": height,
                "bytes": blob,
            })

    return images


def collect_image_pages(structured: dict):
    """Return source pages that contain questions marked image_required."""
    pages = set()
    for sec in structured.get("sections", []):
        for q in sec.get("questions", []):
            if q.get("image_required") and q.get("source_page"):
                try:
                    pages.add(int(q.get("source_page")))
                except Exception:
                    pass
    return sorted(pages)



def extract_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def extract_answer_text(uploaded) -> str:
    if uploaded is None:
        return ""
    data = uploaded.getvalue()
    suffix = Path(uploaded.name).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(data)
    if suffix == ".docx":
        return extract_docx_text(data)
    return data.decode("utf-8", errors="ignore")



def parse_answer_key(answer_text: str) -> dict:
    """
    Parse simple answer keys locally, e.g.
    (1) A
    (2) B
    3. D
    4 B
    (5) A, C
    No Gemini call is needed.
    """
    answer_map = {}
    if not answer_text:
        return answer_map

    for raw_line in answer_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # Question number + one or more English answer letters.
        m = re.match(
            r"^\s*[\(\[]?\s*(\d+)\s*[\)\]\.\:\-]*\s*"
            r"([A-Za-z](?:\s*[,/;&+]\s*[A-Za-z])*)\s*$",
            line,
        )
        if not m:
            continue

        qno = str(int(m.group(1)))
        letters = re.findall(r"[A-Za-z]", m.group(2).upper())
        if letters:
            answer_map[qno] = list(dict.fromkeys(letters))

    return answer_map


def normalize_question_number(value, fallback=None):
    """Normalize '(12)', '12.', 'Q12' etc. to '12' for local answer matching."""
    text = str(value or "").strip()
    m = re.search(r"\d+", text)
    if m:
        return str(int(m.group(0)))
    return str(fallback) if fallback is not None else ""


def apply_local_answer_key(structured: dict, answer_map: dict):
    """
    Attach answers locally by question number.
    Only explicit letters from the uploaded answer file are used.
    """
    seq = 1
    matched = 0

    for sec in structured.get("sections", []):
        for q in sec.get("questions", []):
            qno = normalize_question_number(q.get("number"), fallback=seq)
            letters = answer_map.get(qno, [])
            q["answer"] = letters
            if letters:
                matched += 1
            seq += 1

    return structured, matched



def clean_json_text(s: str) -> str:
    s = s.strip()
    s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.I)
    s = re.sub(r"\s*```$", "", s)
    start = s.find("{")
    end = s.rfind("}")
    if start >= 0 and end > start:
        s = s[start:end+1]
    return s


def image_inventory_text(images) -> str:
    if not images:
        return "(No extractable raster images detected.)"
    rows = []
    for im in images:
        rows.append(
            f"{im['id']} | page={im['page']} | size={im['width']}x{im['height']} | file={im['filename']}"
        )
    return "\n".join(rows)


def build_prompt(form_text: str) -> str:
    """Compact prompt for speed and low token usage."""
    return f"""
Parse this Google Forms PDF text into JSON for Microsoft Forms Quick Import.

ONLY 3 TYPES:
c = choice
m = multiple_answers
t = open_text

MAP:
- Multiple choice / Dropdown / Linear scale -> c
- Grid / Likert / Matching -> split each row -> c
- Checkboxes / Checkbox grid -> m; split checkbox-grid rows
- Short answer / Paragraph / unsafe -> t

RULES:
- Preserve original wording, order and options.
- If visible options exist, keep them.
- Never invent missing text or answers.
- For image-dependent questions: img=true and p=source PAGE number when clear.
- JSON only.

SCHEMA:
{{
  "title": "",
  "sections": [
    {{
      "title": "",
      "questions": [
        {{
          "n": "",
          "q": "",
          "t": "c|m|t",
          "o": [],
          "img": false,
          "p": null,
          "r": false
        }}
      ]
    }}
  ]
}}

FORM:
{form_text}
"""

def expand_compact_structure(raw_structured: dict) -> dict:
    """Convert compact Gemini JSON into the app's normal internal structure."""
    result = {
        "title": raw_structured.get("title", ""),
        "description": raw_structured.get("description", ""),
        "sections": [],
    }

    for sec in raw_structured.get("sections", []):
        new_sec = {
            "title": sec.get("title", ""),
            "description": sec.get("description", ""),
            "questions": [],
        }

        for q in sec.get("questions", []):
            t = str(q.get("t", "t")).strip().lower()
            mapped = {
                "c": "choice",
                "m": "multiple_answers",
                "t": "open_text",
                "choice": "choice",
                "multiple_answers": "multiple_answers",
                "open_text": "open_text",
            }.get(t, "open_text")

            new_sec["questions"].append({
                "number": str(q.get("n", "") or ""),
                "question": str(q.get("q", "") or ""),
                "output_type": mapped,
                "options": list(q.get("o", []) or []),
                "answer": [],
                "required": False,
                "image_required": bool(q.get("img", False)),
                "source_page": q.get("p"),
                "image_refs": [],
                "review_required": bool(q.get("r", False)),
            })

        result["sections"].append(new_sec)

    return result


def normalize_for_quick_import(structured: dict):
    """Enforce only 3 output types: choice, multiple_answers, open_text."""
    for sec in structured.get("sections", []):
        for q in sec.get("questions", []):
            qtype = str(q.get("output_type", "open_text") or "open_text").strip().lower()
            opts = [str(x).strip() for x in (q.get("options") or []) if str(x).strip()]

            if qtype not in ("choice", "multiple_answers", "open_text"):
                qtype = "open_text"
                q["review_required"] = True

            if qtype in ("choice", "multiple_answers"):
                if len(opts) >= 2:
                    q["output_type"] = qtype
                    q["options"] = opts
                else:
                    q["output_type"] = "open_text"
                    q["options"] = []
                    q["review_required"] = True
            else:
                q["output_type"] = "open_text"
                q["options"] = []

    return structured

def classify_gemini_error(exc: Exception) -> str:
    msg = str(exc).lower()

    if ("service_disabled" in msg or
        "gemini api has not been used in project" in msg or
        ("generativelanguage.googleapis.com" in msg and "disabled" in msg)):
        return "service_disabled"

    if ("api_key_invalid" in msg or "api key not valid" in msg or
        "invalid api key" in msg or "unauthenticated" in msg or "401" in msg):
        return "invalid_key"

    if ("permission_denied" in msg or "caller does not have permission" in msg or "403" in msg):
        return "permission_denied"

    if ("429" in msg or "resource_exhausted" in msg or "quota_exceeded" in msg or
        "too_many_requests" in msg or "rate limit" in msg or "quota exceeded" in msg):
        return "rate_limit"

    if ("503" in msg or "service_unavailable" in msg or "high demand" in msg or
        "temporarily unavailable" in msg or "service unavailable" in msg):
        return "service_busy"

    if ("404" in msg or "not_found" in msg or "no longer available" in msg or
        "model not available" in msg):
        return "model_unavailable"

    return "fatal"


def call_gemini_with_retry(api_key: str, prompt: str, status_box=None):
    """Retry once only for 429 or 503; keep the real error category."""
    client = genai.Client(api_key=api_key)
    model_name = DEFAULT_MODEL

    for attempt in range(2):
        try:
            if status_box:
                if attempt == 0:
                    status_box.write(f"正在连接 Gemini：{model_name}")
                else:
                    status_box.write("正在进行最后一次自动重试…")

            response = client.models.generate_content(
                model=model_name,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                ),
            )
            return response, model_name

        except Exception as exc:
            kind = classify_gemini_error(exc)

            if kind in ("rate_limit", "service_busy") and attempt == 0:
                if status_box:
                    if kind == "rate_limit":
                        status_box.write("Free API 使用量暂时达到限制（429），3 秒后自动重试一次…")
                    else:
                        status_box.write("Gemini 服务目前繁忙（503），3 秒后自动重试一次…")
                time.sleep(3)
                continue

            # 403 / SERVICE_DISABLED / invalid key / 404 and second 429/503:
            # stop immediately and let the UI show the exact category.
            raise


def show_clear_api_error(exc: Exception):
    kind = classify_gemini_error(exc)

    if kind == "rate_limit":
        st.error("🟡 Free API 使用额度或请求频率暂时达到限制（429）。Mini App 已自动重试一次，请稍后再按 Convert。")
        st.caption("这通常与当前 Google Cloud Project 的 RPM / TPM / RPD 限制有关，不是 PDF 格式错误。")
    elif kind == "service_busy":
        st.error("🟠 Gemini 服务目前繁忙（503 High Demand）。Mini App 已自动重试一次，请稍后再按 Convert。")
        st.caption("这是 Gemini 服务器容量问题，不代表 PDF 或 API Key 损坏。")
    elif kind == "service_disabled":
        st.error("🔴 Gemini API 尚未启用（SERVICE_DISABLED）。请先为这个 API Key 所属 Project 启用 Gemini API。")
    elif kind == "permission_denied":
        st.error("🔴 API Key / Google Cloud Project 没有 Gemini 内容生成权限（403 PERMISSION_DENIED）。")
    elif kind == "invalid_key":
        st.error("🔴 Gemini API Key 无效或认证失败，请重新检查 API Key。")
    elif kind == "model_unavailable":
        st.error(f"🔴 当前 Gemini Model `{DEFAULT_MODEL}` 不可用，请更新 GEMINI_MODEL 后再试。")
    else:
        st.error("转换失败：Gemini API 返回未识别的错误。")
        with st.expander("查看技术错误（排查时才打开）"):
            st.code(str(exc))


def attach_page_image_candidates(structured: dict, images):
    """
    Map extracted PDF images to questions by source page locally.
    Gemini does not receive image inventory, reducing input tokens.
    """
    by_page = defaultdict(list)
    for im in images:
        by_page[im.get("page")].append(im.get("id"))

    for sec in structured.get("sections", []):
        for q in sec.get("questions", []):
            if q.get("image_required"):
                page = q.get("source_page")
                refs = [x for x in by_page.get(page, []) if x]
                q["image_refs"] = refs
                if len(refs) != 1:
                    q["review_required"] = True
            else:
                q["image_refs"] = []

    return structured


def extract_usage_tokens(response):
    """Read Gemini usage metadata when the SDK returns it."""
    usage = getattr(response, "usage_metadata", None)
    if usage is None:
        return {"input": None, "output": None, "total": None}

    def get_value(*names):
        for name in names:
            value = getattr(usage, name, None)
            if value is not None:
                return int(value)
        return None

    return {
        "input": get_value("prompt_token_count", "input_token_count"),
        "output": get_value("candidates_token_count", "output_token_count"),
        "total": get_value("total_token_count"),
    }



def build_image_lookup(images):
    return {im["id"]: im for im in images}


def option_label(index: int) -> str:
    """Excel-style letters: A..Z, AA..AZ..."""
    n = index + 1
    letters = ""
    while n:
        n, rem = divmod(n - 1, 26)
        letters = chr(65 + rem) + letters
    return letters


def add_image_to_docx(doc, image_bytes, max_width_inches=5.8):
    """Insert image with conservative sizing. Returns True on success."""
    try:
        stream = io.BytesIO(image_bytes)
        doc.add_picture(stream, width=Inches(max_width_inches))
        return True
    except Exception:
        try:
            stream = io.BytesIO(image_bytes)
            doc.add_picture(stream)
            return True
        except Exception:
            return False



def normalize_answer_letters(answer_values, options):
    """Return explicit answers as English option letters; never guess."""
    if answer_values is None:
        return []
    if isinstance(answer_values, str):
        answer_values = [answer_values]

    opts = [str(x).strip() for x in (options or [])]
    result = []

    for raw in answer_values:
        value = str(raw).strip()
        if not value:
            continue

        m = re.fullmatch(r"\(?\s*([A-Za-z]{1,3})\s*\)?[.\s]*", value)
        if m:
            letters = m.group(1).upper()
            if opts:
                idx = 0
                for ch in letters:
                    idx = idx * 26 + (ord(ch) - 64)
                idx -= 1
                if 0 <= idx < len(opts) and letters not in result:
                    result.append(letters)
            elif letters not in result:
                result.append(letters)
            continue

        for idx, opt in enumerate(opts):
            if value == opt:
                letter = option_label(idx)
                if letter not in result:
                    result.append(letter)
                break

    return result


def make_docx(structured: dict, quiz_mode: bool, images):
    """
    Microsoft Forms Quick Import guidance format:
    - Multiple choice: question + vertically stacked A./B./C. options
    - Open text: question only, no answer lines/placeholders
    - Clear blank paragraph between questions
    - NO images/figures embedded in Quick Import Word
    """
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    mapping_rows = []
    embedded_ids = set()  # Intentionally empty for Quick Import Word.

    title = (structured.get("title") or "Microsoft Forms Import").strip()
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(16)

    desc = (structured.get("description") or "").strip()
    if desc:
        doc.add_paragraph(desc)

    global_q_no = 1

    for sec in structured.get("sections", []):
        sec_title = (sec.get("title") or "").strip()
        sec_desc = (sec.get("description") or "").strip()

        # Keep section text simple and vertical.
        if sec_title:
            p = doc.add_paragraph()
            r = p.add_run(sec_title)
            r.bold = True

        if sec_desc:
            doc.add_paragraph(sec_desc)

        for q in sec.get("questions", []):
            original_num = str(q.get("number") or "").strip()
            text = str(q.get("question") or "").strip()
            qtype = str(q.get("output_type") or "open_text").strip()
            image_required = bool(q.get("image_required", False))
            source_page = q.get("source_page")
            image_refs = list(q.get("image_refs") or [])

            # Prefer original number if present; otherwise use sequential numbering.
            number_text = original_num if original_num else str(global_q_no)
            if number_text.endswith("."):
                question_line = f"{number_text} {text}".strip()
            else:
                question_line = f"{number_text}. {text}".strip()

            # Plain vertical question line.
            doc.add_paragraph(question_line)

            opts = [str(x).strip() for x in (q.get("options") or []) if str(x).strip()]

            # Microsoft Forms recommended Multiple choice layout.
            if qtype in ("choice", "multiple_answers"):
                for idx, opt in enumerate(opts):
                    doc.add_paragraph(f"{option_label(idx)}. {opt}")

            # Open text: no answer line, no placeholder, no technical marker.
            elif qtype == "open_text":
                pass

            # Do not embed images in Quick Import Word.
            if image_required:
                mapping_rows.append({
                    "question": question_line,
                    "page": source_page,
                    "image_refs": image_refs,
                    "status": "saved_separately_for_manual_insert",
                })

            # Quiz mode: write explicit answers using English option letters.
            # Example: Answer: A  /  Answer: A, C
            if quiz_mode and qtype in ("choice", "multiple_answers"):
                answer_letters = normalize_answer_letters(q.get("answer") or [], opts)
                if answer_letters:
                    p_ans = doc.add_paragraph()
                    r_ans = p_ans.add_run("Answer: " + ", ".join(answer_letters))
                    r_ans.bold = True

            # Clear separation between questions.
            doc.add_paragraph("")
            global_q_no += 1

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue(), mapping_rows, embedded_ids


def safe_stem(name: str) -> str:
    stem = Path(name).stem
    return re.sub(r'[\\/:*?"<>|]+', "_", stem)


def safe_docx_filename(name: str) -> str:
    return f"{safe_stem(name)}_Microsoft_Forms_Import_V6_10.docx"


def make_mapping_text(mapping_rows, images):
    lines = [
        "Coach Winnie – Forms Converter V6.10",
        "Image Mapping Report",
        "",
        "Important: Quick Import Word intentionally contains NO images, following Microsoft Forms import guidance.",
        "Use this report and the images folder to add images manually after Quick Import.",
        "",
    ]

    if mapping_rows:
        for i, row in enumerate(mapping_rows, start=1):
            refs = ", ".join(row["image_refs"]) if row["image_refs"] else "(none)"
            lines.append(f"{i}. Question: {row['question']}")
            lines.append(f"   PDF page: {row.get('page') or '(unknown)'}")
            lines.append(f"   Image refs: {refs}")
            lines.append(f"   Status: {row['status']}")
            lines.append("")
    else:
        lines.append("No image-dependent questions were identified.")
        lines.append("")

    if images:
        lines.append("Extracted image inventory:")
        for im in images:
            lines.append(
                f"- {im['id']} -> images/{im['filename']} "
                f"(page {im['page']}, {im['width']}x{im['height']})"
            )
    else:
        lines.append("No extractable raster images were detected in the PDF.")

    return "\n".join(lines)


def make_bundle_zip(pdf_name: str, docx_bytes: bytes, images, mapping_text: str) -> bytes:
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(safe_docx_filename(pdf_name), docx_bytes)
        zf.writestr("image_mapping.txt", mapping_text)
        for im in images:
            zf.writestr(f"images/{im['filename']}", im["bytes"])
    return bio.getvalue()


def build_conversion_report(structured: dict, images, embedded_ids):
    questions = [
        q for sec in structured.get("sections", [])
        for q in sec.get("questions", [])
    ]
    output_types = Counter(q.get("output_type", "open_text") for q in questions)

    return {
        "total": len(questions),
        "output_types": output_types,
        "review_count": sum(1 for q in questions if q.get("review_required")),
        "image_questions": sum(1 for q in questions if q.get("image_required")),
        "extracted_images": len(images),
        "embedded_images": len(embedded_ids),
    }


if st.button("✨ Convert to Microsoft Forms Word", type="primary", use_container_width=True):
    if not pdf_file:
        st.error("请先上传 Google Forms PDF。")
        st.stop()

    api_key = api_key_input.strip()
    if not api_key:
        st.error("请先输入你自己的 Gemini API Key。")
        st.stop()

    try:
        with st.status("正在读取并转换表单…", expanded=True) as status:
            st.write("读取 PDF 文字")
            pdf_bytes = pdf_file.getvalue()
            raw_form_text = extract_pdf_text(pdf_bytes)

            st.write("清理 Google Forms UI 文字")
            form_text = clean_form_text_for_gemini(raw_form_text)

            st.write("读取答案文件" if answer_file else "没有答案文件，将不会推测答案")
            answer_text = extract_answer_text(answer_file)
            answer_map = parse_answer_key(answer_text)

            quiz_mode = output_mode.startswith("Quiz")
            prompt = build_prompt(form_text)

            st.write("Fast + Low Token：Gemini 只识别题目结构与 3 种题型")
            response, model_used = call_gemini_with_retry(
                api_key, prompt, status
            )
            st.write(f"Gemini 连接成功：{model_used}")
            raw = response.text or ""
            compact_structured = json.loads(clean_json_text(raw))
            structured = expand_compact_structure(compact_structured)
            structured = normalize_for_quick_import(structured)

            # Lazy image extraction: only scan pages actually marked as image questions.
            image_pages = collect_image_pages(structured)
            if image_pages:
                st.write(f"只提取图片题所在页面：{', '.join(map(str, image_pages))}")
                images = extract_pdf_images_for_pages(pdf_bytes, image_pages)
            else:
                images = []
                st.write("未检测到图片题，跳过图片提取")

            structured = attach_page_image_candidates(structured, images)
            matched_answers = 0
            if quiz_mode and answer_map:
                structured, matched_answers = apply_local_answer_key(structured, answer_map)
            else:
                structured, _ = apply_local_answer_key(structured, {})

            usage_tokens = extract_usage_tokens(response)

            st.write("Python 正在处理答案、图片对应与 Word 排版")
            docx_bytes, mapping_rows, embedded_ids = make_docx(
                structured, quiz_mode, images
            )
            mapping_text = make_mapping_text(mapping_rows, images)
            bundle_bytes = make_bundle_zip(
                pdf_file.name, docx_bytes, images, mapping_text
            )
            report = build_conversion_report(structured, images, embedded_ids)
            status.update(label="转换完成 ✅", state="complete")

        st.success(f"完成：共生成 {report['total']} 个 Quick Import 兼容题目。")

        st.subheader("📊 转换报告")
        c1, c2, c3 = st.columns(3)
        c1.metric("Choice", report["output_types"].get("choice", 0))
        c2.metric("Multiple-answer Choice", report["output_types"].get("multiple_answers", 0))
        c3.metric("Open text", report["output_types"].get("open_text", 0))

        c4, c5, c6 = st.columns(3)
        c4.metric("已提取图片", report["extracted_images"])
        c5.metric("图片题", report["image_questions"])
        c6.metric("Quick Import 内图片", report["embedded_images"])

        st.subheader("⚡ Fast + Low Token Mode")
        t1, t2, t3 = st.columns(3)
        t1.metric("Input Tokens", usage_tokens["input"] if usage_tokens["input"] is not None else "—")
        t2.metric("Output Tokens", usage_tokens["output"] if usage_tokens["output"] is not None else "—")
        t3.metric("Total Tokens", usage_tokens["total"] if usage_tokens["total"] is not None else "—")

        if quiz_mode and answer_file:
            st.caption(
                f"✅ 答案文件由 Python 本地配对：读取 {len(answer_map)} 个答案，"
                f"成功对应 {matched_answers} 题；答案内容没有发送给 Gemini。"
            )

        notes = []
        multi_count = report["output_types"].get("multiple_answers", 0)
        if multi_count:
            notes.append(
                f"Multiple-answer Choice：{multi_count} 题。"
                "Microsoft Forms Quick Import 后，请为这些题开启 Multiple answers。"
            )
        if report["image_questions"]:
            notes.append(
                "图片题：为了符合 Microsoft Forms Import Guidance，Quick Import Word 不放图片。"
                "原图会保存在 ZIP 的 images 文件夹，请按 image_mapping.txt 在导入后补回。"
            )
        if report["review_count"]:
            notes.append(
                f"人工检查：{report['review_count']} 题因选项不足、结构不清或图片配对不确定，"
                "已尽量安全转换，建议导入后核对。"
            )

        if notes:
            st.warning("\n\n".join(notes))
        else:
            st.info("转换完成，没有检测到需要额外人工检查的题目。")


        st.caption("✅ Quick Import Word 已按 Microsoft Forms Import Guidance 排版；Quiz 模式有答案文件时，选择题会加入英文格式 Answer: A。")

        st.download_button(
            "⬇️ Download Microsoft Forms Word",
            data=docx_bytes,
            file_name=safe_docx_filename(pdf_file.name),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.download_button(
            "📦 Download Word + Extracted Images ZIP",
            data=bundle_bytes,
            file_name=f"{safe_stem(pdf_file.name)}_Microsoft_Forms_V6_10_Bundle.zip",
            mime="application/zip",
            use_container_width=True,
        )

        with st.expander("查看 Image Mapping Report"):
            st.code(mapping_text, language=None)

        st.caption(
            "建议流程：Microsoft Forms → Quick Import → 导入 Word → "
            "检查 Choice / Multiple-answer Choice / Open text → "
            "如果图片未自动导入，打开 ZIP 的 images 文件夹，根据 image_mapping.txt 加回对应题目。"
        )

    except json.JSONDecodeError:
        st.error("Gemini 输出无法解析为结构化数据，请重试。若 PDF 很复杂，可分成较小部分转换。")
    except Exception as e:
        show_clear_api_error(e)
        if "API_KEY_INVALID" in msg or "API key not valid" in msg or "INVALID_ARGUMENT" in msg:
            st.error("Gemini API Key 无效或请求设置不正确。请到 Google AI Studio 检查 API Key 后再试。")
        elif "404" in msg or "NOT_FOUND" in msg or "no longer available" in msg:
            st.error(
                "当前 Gemini 模型对此 API Key / project 不可用。"
                "请检查 Google AI Studio 中该 project 可使用的模型。"
            )
        elif (
            "RESOURCE_EXHAUSTED" in msg
            or "429" in msg
            or "503" in msg
            or "UNAVAILABLE" in msg
            or "服务目前繁忙" in msg
        ):
            st.error(
                "Gemini 服务目前繁忙或暂时达到 Free API 使用限制。"
                "Mini App 已自动重试一次，请稍后再按 Convert。"
            )
        else:
            st.error(f"转换失败：{e}")

st.divider()
st.markdown("""
<div class="small">
<strong>Microsoft Forms 导入：</strong>
Microsoft Forms → Quick Import → Upload from this device → 选择生成的 Word → Form / Quiz。<br>
<strong>V6.10 Clear API Status：</strong>
Python 会先清理 PDF UI 文字；Gemini 使用 Compact JSON，只负责题目解析与 3 种题型分类；图片仅在检测到图片题时才按相关页面 Lazy Extraction。<br>Quiz 模式若上传答案文件，选择题答案仍以英文选项字母输出，例如 <strong>Answer: A</strong>；没有答案的题不会推测。<br>
<strong>重要限制：</strong>
Microsoft Forms Quick Import 不保证把 Word 内图片自动转换成题目图片；若图片没有带入，请使用 ZIP 内原图手动补回。<br>
<strong>Privacy：</strong>
上传内容会使用你本次输入的 Gemini API Key 发送到 Google Gemini API 进行转换；
本 App 不会将 API Key 写入 GitHub、文件或数据库。
</div>
""", unsafe_allow_html=True)
